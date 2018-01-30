# -*- coding: utf-8 -*-


#请在python3下的环境编译
#安装依赖pip3 install python-docx



from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os

document = Document()

document.add_heading('HHU ACM Template', 0)

document.add_page_break()

def Handle_Categeory(name) :
    path = "./" + name
    files = os.listdir(path)
    index = 0
    if(len(files) == 0) :
        return
    head = document.add_heading('-------' + name + '-------', level=1)
    head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for file in files:
        if not os.path.isdir(file) and os.path.splitext(file)[1] == '.md':
            # print(os.path.splitext(file)[0])
            index = index + 1
            document.add_heading(str(index) + '--' + os.path.splitext(file)[0] , level=2)
            with open(path + '/' + file , mode='r', encoding='utf-8') as file:
                content = str(file.read()).strip();
                para = document.add_paragraph(content, style='ListParagraph')
                # para_format = para.paragraph_format
                # para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para_run = para.add_run()
                font = para_run.font
                font.size = Pt(10)
    document.add_page_break()
categeories = ('Datastructure', 'Dynamic', 'Geometry', 'Graph', 'Math', 'Others', 'String')
for categeory in categeories:
    Handle_Categeory(categeory)
document.save('./PDF/ACM.docx')